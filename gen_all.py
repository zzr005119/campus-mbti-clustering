# -*- coding: utf-8 -*-
"""Generate both course reports"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

BASE = r"D:\Xx\大二下\机器学习\大作业"
SIDS = "2024117715\n2024117720\n2024117719"
SNAMES = "赵子荣\n李海东\n章振豪"
CLASS = "计算机科学与技术2022级X班"
TEACHER = "曾梦"

print("Base init OK")
CLUSTER_DATA = {
    0: {"name": "高消费数码型", "count": 12, "budget": 4.08, "food": 3.5,
        "social": 2.58, "fashion": 1.67, "study": 1.67, "digital": 3.17,
        "hobby": 2.08, "mbti": "XNTP",
        "desc": "预算充裕，热爱数码与游戏，兴趣投入高，是校园里的装备党和游戏达人。"},
    1: {"name": "节俭生存型", "count": 23, "budget": 2.13, "food": 2.04,
        "social": 1.43, "fashion": 1.13, "study": 1.7, "digital": 2.17,
        "hobby": 1.57, "mbti": "ISFP",
        "desc": "精打细算的务实派，各项消费节制，偏好高性价比的生活方式。"},
    2: {"name": "社交体验型", "count": 14, "budget": 3.07, "food": 2.57,
        "social": 2.64, "fashion": 2.21, "study": 1.57, "digital": 1.79,
        "hobby": 1.14, "mbti": "XSFP",
        "desc": "社交达人，热衷聚餐和娱乐活动，注重外在形象，享受和朋友在一起的时光。"},
    3: {"name": "自我提升型", "count": 13, "budget": 3.15, "food": 2.69,
        "social": 2.46, "fashion": 2.77, "study": 3.0, "digital": 2.08,
        "hobby": 2.38, "mbti": "XNFP",
        "desc": "注重自我成长，在学习发展和兴趣培养上投入最多，是校园里的充电型选手。"},
}

def add_personnel(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run("人员分配：")
    r.bold = True; r.font.name = "宋体"; r.font.size = Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run("本项目共三位成员协作完成，具体分工如下：").font.size = Pt(12)
    assigns = [
        ("（1）赵子荣（2024117715）：", "负责项目总体设计与技术架构、核心算法实现（K-Means聚类与随机森林分类器）、LOOCV交叉验证与SMOTENC过采样策略设计、Bootstrap置信区间评估、模型训练与调优、可视化图表生成、Streamlit与Flask双版本推荐系统开发、答辩PPT制作与展示。同时统筹两门课程报告的撰写框架与内容整合。"),
        ("（2）李海东（2024117720）：", "负责前期文献调研与MBTI性格理论梳理、问卷设计与问卷星平台发放、原始数据收集与初步整理、数据清洗规则设计。参与课程报告相关理论与算法基础章节的撰写与参考文献整理，协助实验报告中的实验环境与技术准备章节编写。"),
        ("（3）章振豪（2024117719）：", "负责数据可视化方案设计（雷达图、热力图、PCA散点图等）、Web前端界面美化与交互优化、用户测试与系统体验评估。参与实验报告实验结果与可视化分析章节的图表演示与分析解读，协助课程报告中的应用场景与效果分析章节撰写。"),
    ]
    for nm, desc in assigns:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        rn = p.add_run(nm); rn.bold = True; rn.font.size = Pt(12)
        p.add_run(desc).font.size = Pt(12)

def add_abstract(doc, atype="course"):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run("摘要："); r.bold = True; r.font.name = "宋体"; r.font.size = Pt(12)
    if atype == "course":
        t = "本研究以62份真实大学生问卷数据为基础，综合运用K-Means无监督聚类算法与随机森林分类器，探索校园消费行为与MBTI人格特质之间的深层关联。首先通过肘部法则与轮廓系数确定最优聚类数K=4，识别出高消费数码型、节俭生存型、社交体验型、自我提升型四类消费群体。进而以12维消费与人口学特征为输入，分别预测MBTI四个维度的倾向，采用留一法交叉验证（LOOCV）与SMOTENC过采样应对小样本与类别不平衡挑战，并引入Bootstrap 95%置信区间与Dummy Classifier基线进行严格评估。实验结果表明：J/P生活风格维度预测效果最佳（Acc=58.06%, F1=0.580, AUC=0.647），S/N信息处理维度次之（Acc=61.29%），T/F决策标准维度最难预测（Acc=45.16%）。特征重要性分析揭示了社交消费与外倾倾向、服饰消费与情感决策、看重因素与计划性等有趣的消费-人格映射关系。最终基于聚类与分类模型，设计并实现了融合消费相似度、MBTI互补性与场景偏好的校园搭子推荐系统。"
    else:
        t = "本实验基于62份校园问卷数据，使用Python及scikit-learn库完整实现了K-Means消费群体聚类、随机森林MBTI性格预测与Streamlit/Flask双版本搭子推荐系统。实验涵盖数据清洗（剔除填写时间<30秒及消费五维全同的无效样本）、StandardScaler标准化、肘部法则与轮廓系数联合确定K=4、LOOCV留一法交叉验证、SMOTENC过采样处理类别不平衡、Bootstrap 95%置信区间评估、Dummy Classifier基线对比等关键环节。通过matplotlib与seaborn生成了肘部法则图、雷达图、热力图、PCA散点图、混淆矩阵、ROC曲线、指标对比柱状图和特征重要性条形图共8类可视化成果。最终交付了完整的Python代码包（约550行核心代码）、训练好的模型文件（K-Means + 4个RF分类器 + 5个Scaler）以及可交互的Web推荐系统原型。"
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(t).font.size = Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run("关键词："); r.bold = True; r.font.size = Pt(12)
    kw = "K-Means聚类；随机森林；MBTI人格预测；留一法交叉验证；搭子推荐系统" if atype == "course" else "K-Means聚类；随机森林；MBTI人格预测；SMOTENC过采样；校园搭子推荐系统"
    p.add_run(kw).font.size = Pt(12)

def setup_doc():
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "宋体"; s.font.size = Pt(12)
    s.paragraph_format.line_spacing = 1.5
    s.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return doc

def add_cover(doc, cname, ccode, title, is_exp=False):
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本 科 生 课 程 报 告 封 面")
    r.font.size = Pt(22); r.font.name = "黑体"; r.bold = True
    doc.add_paragraph()
    tbl = doc.add_table(rows=7, cols=2); tbl.autofit = True
    data = [
        ("课程名称：", cname), ("课程编码：", ccode), ("报告题目：", title),
        ("学    号：", SIDS), ("姓    名：", SNAMES),
        ("班    级：", CLASS), ("授课教师：", TEACHER),
    ]
    for i, (lb, vl) in enumerate(data):
        row = tbl.rows[i]
        p0 = row.cells[0].paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0 = p0.add_run(lb); r0.font.size = Pt(14); r0.font.name = "宋体"
        r0.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        p1 = row.cells[1].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(vl)
        r1.font.size = Pt(11) if i == 2 else Pt(14); r1.font.name = "宋体"
        r1.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if i == 2: r1.bold = True
    tblP = tbl._tbl.tblPr
    if tblP is None:
        tblP = parse_xml('<w:tblPr ' + nsdecls("w") + '></w:tblPr>')
        tbl._tbl.insert(0, tblP)
    bdrs = parse_xml(
        '<w:tblBorders ' + nsdecls("w") + '>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>')
    tblP.append(bdrs)
    doc.add_paragraph(); doc.add_paragraph()
    if is_exp: _score_exp(doc)
    else: _score_course(doc)
    doc.add_page_break()
    return doc

def _score_course(doc):
    tbl = doc.add_table(rows=5, cols=9); tbl.autofit = True
    hdrs = ["评分标准及分值", "", "选题契合\n（分值20）", "", "理论阐述\n（分值30）", "",
            "应用分析\n（分值30）", "", "报告规范\n（分值20）"]
    for j, h in enumerate(hdrs):
        p = tbl.rows[0].cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.font.size = Pt(9); r.bold = True
    tbl.rows[0].cells[0].merge(tbl.rows[0].cells[1])
    tbl.rows[0].cells[2].merge(tbl.rows[0].cells[3])
    tbl.rows[0].cells[4].merge(tbl.rows[0].cells[5])
    tbl.rows[0].cells[6].merge(tbl.rows[0].cells[7])
    p = tbl.rows[1].cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("评分").font.size = Pt(10)
    tbl.rows[2].cells[0].merge(tbl.rows[2].cells[8])
    tbl.rows[2].cells[0].paragraphs[0].add_run("注：以上为参考标准，授课教师根据需要对评分标准进行调整").font.size = Pt(9)
    tbl.rows[3].cells[0].paragraphs[0].add_run("评语：").font.size = Pt(10)
    tbl.rows[4].cells[0].paragraphs[0].add_run("总 评 分").font.size = Pt(10)
    p = tbl.rows[4].cells[4].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("评阅教师：").font.size = Pt(10)
    p = tbl.rows[4].cells[7].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("评阅时间\n2026年 月 日").font.size = Pt(10)

def _score_exp(doc):
    tbl = doc.add_table(rows=5, cols=8); tbl.autofit = True
    hdrs = ["评分标准及分值", "", "选题契合\n（分值15）", "报告内容与代码实现\n（分值40）", "",
            "实验分析\n（分值30）", "", "报告规范\n（分值15）"]
    for j, h in enumerate(hdrs):
        p = tbl.rows[0].cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.font.size = Pt(9); r.bold = True
    tbl.rows[0].cells[0].merge(tbl.rows[0].cells[1])
    tbl.rows[0].cells[3].merge(tbl.rows[0].cells[4])
    tbl.rows[0].cells[5].merge(tbl.rows[0].cells[6])
    p = tbl.rows[1].cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("评分").font.size = Pt(10)
    tbl.rows[2].cells[0].merge(tbl.rows[2].cells[7])
    tbl.rows[2].cells[0].paragraphs[0].add_run("注：以上为参考标准，授课教师根据需要对评分标准进行调整").font.size = Pt(9)
    tbl.rows[3].cells[0].paragraphs[0].add_run("评语：").font.size = Pt(10)
    tbl.rows[4].cells[0].paragraphs[0].add_run("总 评 分").font.size = Pt(10)
    p = tbl.rows[4].cells[3].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("评阅教师：").font.size = Pt(10)
    p = tbl.rows[4].cells[6].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("评阅时间\n2026年 月 日").font.size = Pt(10)

def add_h(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "黑体"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return h

def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(text).font.size = Pt(12)
    return p

def add_pb(doc):
    doc.add_page_break()

print("Helper functions defined OK")
def gen_course_report():
    doc = setup_doc()
    title = "基于K-Means聚类与随机森林的\n校园消费群体画像与MBTI人格预测研究"
    doc = add_cover(doc, "机器学习", "2351302", title, is_exp=False)
    add_personnel(doc)
    add_pb(doc); add_abstract(doc, "course"); add_pb(doc)

    add_h(doc, "一、选题背景与研究意义", 1)
    add_h(doc, "1.1 选题背景", 2)
    add_p(doc, "在大学校园中，消费行为和人格特质是影响学生社交匹配的两个核心维度。消费水平与消费结构直接影响学生能否玩到一起——预算差异过大的搭子在聚餐、出游等场景中容易产生摩擦和不适感。与此同时，MBTI人格类型被广泛用于解释个体的社交偏好和决策模式：E/I维度决定了一个人是通过社交获取能量还是独处恢复精力；S/N维度影响信息接收方式，进而影响消费决策中对细节的关注程度；T/F维度反映决策标准是偏理性还是偏感性，这在选购商品时表现尤为突出；J/P维度则影响出行计划的灵活性和消费的计划性。然而，目前鲜有研究将微观消费特征与MBTI人格两个维度进行定量结合分析，更缺乏基于数据驱动的校园社交推荐实践。本研究旨在填补这一空白，通过机器学习方法建立消费行为与人格特质之间的映射模型，并据此构建个性化的校园搭子推荐系统。")
    add_h(doc, "1.2 研究意义", 2)
    add_p(doc, "理论价值方面，本研究通过K-Means聚类探索校园消费群体的自然划分，验证消费行为数据在人群细分中的有效性；通过随机森林分类器建立消费到人格的预测模型，丰富了MBTI研究的数据驱动方法论，为心理学与机器学习的交叉研究提供了可复现的实证案例。实践意义方面，基于62份真实问卷数据构建的校园搭子推荐系统，能够自动识别用户的消费群体和人格类型，综合消费相似度、MBTI互补性和场景偏好给出个性化搭子推荐，具有直接的应用落地价值。同时，项目完整覆盖了数据清洗、特征工程、模型训练、交叉验证、可视化分析和Web系统部署的机器学习全流程，对于小样本场景下的机器学习教学实践也具有良好的示范意义。")

    add_pb(doc)
    add_h(doc, "二、相关理论与算法基础", 1)
    add_h(doc, "2.1 K-Means聚类算法", 2)
    add_p(doc, "K-Means是一种基于划分的无监督聚类算法，由MacQueen于1967年提出。其核心目标是将n个样本划分到K个簇中，使得簇内样本的平方欧氏距离之和（Inertia）最小化。设样本集合为X={x1,...,xn}，聚类中心为mu={mu1,...,muK}，则目标函数为：J = sum(i=1 to K) sum(x in Ci) ||x - mui||^2。算法采用EM风格的迭代优化：（1）随机初始化K个聚类中心；（2）E步——将每个样本分配到距离最近的聚类中心所在的簇；（3）M步——重新计算每个簇内样本的均值，更新聚类中心；（4）重复上述步骤直至聚类中心不再变化或达到最大迭代次数。K-Means的核心优势在于算法简洁、收敛速度快（通常10-20轮迭代），且聚类结果具有直观的可解释性。其局限性包括：需要预先指定K值、对初始中心敏感（可能收敛到局部最优）、假设簇为球形分布（欧氏距离在非球形或类别型数据中效果欠佳）。本研究中，K值通过肘部法则与轮廓系数联合确定，以n_init=10的多次随机初始化缓解局部最优问题。")
    add_h(doc, "2.2 随机森林分类算法", 2)
    add_p(doc, "随机森林（Random Forest）是由Breiman于2001年提出的Bagging集成学习方法。其核心思想是通过构建多棵决策树并投票汇总来提高分类精度和泛化能力。每棵决策树的训练过程包含双重随机性：（1）Bootstrap自助采样——从原始训练集中有放回地抽取N个样本构成一棵树的训练集，约有36.8%的样本未被选中（袋外样本OOB），可用于模型评估；（2）随机特征子空间——在每个节点分裂时，从全部M个特征中随机选择m（通常m约等于sqrt(M)）个候选特征，从中选择最优分裂点。随机森林的输出采用多数投票机制：H(x)=argmax_y sum(i=1 to T) I(hi(x)=y)，其中T为树的数量，hi为第i棵决策树。本研究设置n_estimators=100，max_depth=5（控制过拟合），min_samples_leaf=3（确保叶节点有足够样本支持），class_weight=balanced（自动调整类别权重以缓解类别不均衡）。")
    add_h(doc, "2.3 留一法交叉验证（LOOCV）", 2)
    add_p(doc, "留一法（Leave-One-Out Cross-Validation）是K折交叉验证在K=N时的特例：每次取1个样本作为测试集，其余N-1个样本作为训练集，重复N次使得每个样本恰好被测试一次。LOOCV的核心优势在于：（1）训练集利用率极高——每次训练使用了99%以上的数据；（2）评估结果几乎无偏——不同于随机划分可能导致的评估偏差；（3）确定性——每次运行结果完全一致，可复现性强。这些特性使得LOOCV在小样本场景（N=62）下成为最可靠的模型评估策略。然而，LOOCV的局限性在于方差较大（每次仅1个样本评估）和无法分层（在极端不均衡场景下可能低估或高估性能）。本研究在LOOCV基础上进一步引入Bootstrap 95%置信区间，以量化评估的不确定性。")
    add_h(doc, "2.4 SMOTENC过采样", 2)
    add_p(doc, "SMOTE（Synthetic Minority Over-sampling Technique）是由Chawla等人于2002年提出的经典过采样方法：对于每个少数类样本，在其K近邻中随机选择一个邻居，在两者连线上合成新样本：x_new = x + lambda * (x_neighbor - x)，其中lambda在[0,1]区间随机取值。SMOTENC（SMOTE-Nominal Continuous）是SMOTE的扩展变体，专门处理同时包含连续特征和类别特征的混合型数据。在合成新样本时，连续特征采用上述线性插值，类别特征则采用最近邻的众数填充，避免了直接在类别空间插值的不合理性。本研究中，decision_style（消费决策习惯）、priority_factor（选购看重因素）、gender（性别）、income_source（生活费来源）被指定为类别特征，其余8维为连续特征，SMOTENC在每轮LOOCV训练中对训练集进行过采样，以缓解E/I维度中E类仅9份样本的严重不均衡问题。")

    add_pb(doc)
    add_h(doc, "三、理论分析与方案设计", 1)
    add_h(doc, "3.1 核心研究假设", 2)
    add_p(doc, "本研究围绕三个递进层面的假设展开：（1）聚类层面——大学生消费行为存在可被聚类算法识别的自然群体结构，消费五维度（餐饮、社交、服饰、学习、数码）加月预算与兴趣投入构成有效的聚类输入；（2）分类层面——消费行为特征与MBTI人格维度之间存在可被机器学习模型捕获的映射关系，即消费习惯可部分预测MBTI性格，具体表现为：社交花销大指向E倾向、数码花销大指向I倾向、服饰花销大指向F倾向、学习花销大指向J倾向；（3）推荐层面——综合消费相似度、MBTI互补性和场景偏好的多维匹配策略优于单一维度的推荐效果。三层面假设相互关联：聚类结果为分类模型提供群体标签，分类结果为推荐系统提供人格特征，形成感知群体轮廓到预测个体人格再到匹配社交搭子的理论闭环。")
    add_h(doc, "3.2 三阶段递进式方案设计", 2)
    add_p(doc, "阶段一（无监督聚类）：选取5项消费明细（餐饮、社交、服饰、学习、数码）+月生活费+兴趣投入共7维特征，经StandardScaler标准化后以K=2到8遍历K-Means，综合肘部法则拐点（Inertia下降速率变化）与轮廓系数峰值确定最优K。对聚类结果从群体规模、消费均值、MBTI分布三个维度进行群体画像解读。阶段二（有监督分类）：将特征空间扩展至12维（7维消费特征+3维人口学特征+2维消费行为特征），分别以MBTI四维作为预测目标。采用LOOCV留一法评估，SMOTENC处理训练集不均衡，RandomForest（n_estimators=100, max_depth=5, class_weight=balanced）作为基分类器。评估指标体系涵盖Accuracy、Macro F1、Precision、Recall、AUC-ROC、Bootstrap 95%置信区间，并引入DummyClassifier（most_frequent策略）作为基线对比。阶段三（推荐系统）：融合K-Means聚类标签（消费群体匹配）、RF预测的MBTI（人格互补匹配）、Cosine相似度（消费距离匹配）和场景偏好（场景契合度），加权输出TOP-5搭子推荐结果。")
    add_h(doc, "3.3 评估指标选择的理论依据", 2)
    add_p(doc, "在小样本（N=62）且类别不均衡（如E/I维度的9:29:24分布）的场景下，单一Accuracy存在严重误导——Dummy基线（预测多数类）在S/N维度即可获得67.74%的Accuracy。因此本研究采用多指标联合评估：Macro F1对每个类别赋予相同权重，不受样本分布影响，能真实反映模型对所有类别（尤其是少数类）的预测能力；AUC-ROC评估模型对正负样本的排序能力；Bootstrap 95%置信区间（2000次重采样）量化评估的不确定性，避免了在小样本下对指标进行点估计的过度自信。与Dummy基线的Delta差值则揭示了模型是否真正学到了超越随机猜测的信号。")

    add_pb(doc)
    add_h(doc, "四、应用场景与效果分析", 1)
    add_h(doc, "4.1 K-Means聚类适用性分析", 2)
    add_p(doc, "K-Means在本研究中的适用性体现在以下方面：（1）7维特征均为离散等级变量（值域1-4或1-5），经StandardScaler标准化后近似满足欧氏距离对连续型变量的假设；（2）肘部法则显示K=4处Inertia下降速率明显减缓，轮廓系数在K=4时达到局部峰值0.1775，两指标交叉验证了K=4的合理性；（3）聚类结果的业务可解释性极强——四群体在消费结构上差异显著，高消费数码型（n=12, 预算4.08, 数码3.17）、节俭生存型（n=23, 预算2.13, 五维全垫底）、社交体验型（n=14, 社交2.64, 服饰2.21）、自我提升型（n=13, 学习3.0, 服饰2.77, 兴趣2.38），各自对应清晰的校园人物画像，便于在推荐系统中生成直观的群体标签。K-Means的不足在于：假设簇为球形分布，PCA散点图显示部分群体在低维投影中存在重叠，表明实际数据可能存在非球形或非等方差的簇结构，后续可尝试GMM（高斯混合模型）或DBSCAN等对簇形状假设更宽松的聚类算法。")
    add_h(doc, "4.2 随机森林分类效果深度分析", 2)
    add_p(doc, "MBTI四维预测呈现明显的性能分化：J/P生活风格维度表现最优（Acc=58.06%, F1=0.580, AUC=0.647, 高于Dummy基线6.45个百分点），且J、P两类F1均衡（分别为0.57和0.59），说明模型成功捕捉到了消费行为中反映计划性与灵活性差异的信号。选购看重因素（性价比vs颜值vs品牌）成为J/P维度的最强预测特征——这恰好与MBTI理论中J型偏好确定性（性价比）、P型享受过程体验（颜值或品牌）的论述高度吻合。S/N信息处理维度Acc最高（61.29%），但AUC仅0.437且低于Dummy基线，说明高Acc主要来自模型偏向多数类S（F1=0.74）：少数类N仅F1=0.25。社交娱乐消费成为最强特征，支持了偏好社交指向依赖具体感官信息（S）的假设。T/F决策标准维度是唯一低于Dummy基线的（-11.29%），表明仅靠消费行为特征难以有效区分理性思考和情感决策的人——这暗示T/F维度可能更多由价值观念而非消费行为决定。E/I精力恢复维度受限于仅有9份外倾样本，E类F1=0.00，模型完全无法预测外倾类型。但特征重要性排名揭示了一个有趣发现：数码消费重要性排名第2，支持数码花销大指向I倾向（独处恢复精力）的假设。")
    add_h(doc, "4.3 与现有研究的对比", 2)
    add_p(doc, "与典型的人格预测研究相比，本研究在方法论上的创新体现在：（1）以小样本（N=62）为出发点，刻意选择LOOCV而非更常用的K折交叉验证，这是对数据现实约束的正确响应——在小样本下，K折的评估方差远高于LOOCV；（2）不仅报告点估计的Accuracy，引入Bootstrap 95% CI和Dummy基线对比，提供了统计上更完备的评估框架；（3）同时使用SMOTENC处理混合型数据的不均衡问题，避免了传统SMOTE对类别特征的错误处理。在结果层面，J/P维度F1=0.580的表现虽然绝对值不高，但在62份小样本、单题项测量MBTI的约束下已是较好成绩——相比随机猜测（50%）有16%的提升。这为用消费行为预测人格特质这一研究方向提供了初步的实证支持和清晰的改进方向。")
    add_h(doc, "4.4 推荐系统的理论整合", 2)
    add_p(doc, "推荐系统的设计遵循了多维度融合匹配的理论框架：消费相似度（Cosine）确保搭子在预算和消费结构上匹配，避免因消费水平差距过大产生的社交摩擦；MBTI互补性借鉴了人际互补理论——E与I互补（一个带动社交，一个提供深度）、T与F互补（一个决策理性，一个关注感受），而S/N和J/P建议相似（世界观和生活方式的一致性更有利于长期友谊）；场景偏好匹配确保搭子在共同活动中都能获得良好体验。这一多维匹配策略较单一维度的协同过滤具有更强的理论解释力和更好的冷启动适应性——新用户只需填写消费画像即可获得推荐，无需历史行为数据。")

    add_pb(doc)
    add_h(doc, "五、结论与展望", 1)
    add_h(doc, "5.1 主要结论", 2)
    add_p(doc, "本研究以62份校园问卷为数据基础，完整实现了从数据清洗、聚类分析、分类预测到推荐系统部署的机器学习全流程，得出以下主要结论：（1）K-Means聚类成功识别出四类特征鲜明的校园消费群体（高消费数码型、节俭生存型、社交体验型、自我提升型），群体间在消费结构上差异显著，群体画像具有清晰的可解释性。（2）J/P生活风格维度可被消费行为特征最有效地预测（F1=0.580, AUC=0.647），选购看重因素作为最强特征巧妙验证了MBTI理论中J/P维度的核心差异——对计划性和确定性的偏好程度在消费决策中表现突出。（3）S/N维度的Acc较高但AUC偏低，揭示了在不均衡数据下单一Accuracy的误导风险，验证了多指标联合评估的必要性。（4）T/F维度的预测困难（低于Dummy基线）提示了人格的决策维度可能更多由深层价值观而非表层消费行为决定。（5）在62份小样本下，LOOCV + SMOTENC + Bootstrap CI的组合评估框架提供了统计上可靠的性能估计，为小样本机器学习研究提供了方法论参考。")
    add_h(doc, "5.2 局限性与改进方向", 2)
    add_p(doc, "本研究的局限性包括：（1）样本瓶颈——62份样本在机器学习中属于极小样本，E/I维度仅9份外倾样本导致预测完全失效。通过SDV GaussianCopulaSynthesizer生成200-300条合成数据可缓解此问题。（2）MBTI测量简化——每个维度仅用1道题目测量，精度远低于标准93题MBTI量表，自评结果可能受社会期望偏差影响。未来可引入标准化心理量表或行为实验获取更准确的人格标签。（3）特征空间有限——12维特征可能未能覆盖影响人格的消费行为全貌，可引入支付方式偏好、品牌忠诚度、促销敏感度等新特征。（4）聚类方法单一——K-Means对簇形状的假设限制了聚类效果，可尝试GMM（支持非等方差）或谱聚类（支持非球形簇）。（5）模型优化空间——当前RF参数固定（max_depth=5），可通过网格搜索或贝叶斯优化提升性能；亦可尝试XGBoost、LightGBM等梯度提升模型，或构建多任务神经网络同时预测MBTI四维以共享特征表示。")
    add_h(doc, "5.3 未来展望", 2)
    add_p(doc, "未来研究可沿三个方向深化：（1）数据层面——扩大问卷发放至300+样本，引入行为实验和校园一卡通消费流水数据作为客观消费行为的补充，减少自评偏差。（2）模型层面——探索深度学习方案（MLP多任务分类+Autoencoder深度聚类），融合协同过滤和基于内容的推荐算法，并引入主动学习机制在用户使用过程中持续优化匹配策略。（3）应用层面——将推荐系统从单次匹配扩展为社交网络构建工具，通过图神经网络建模用户之间的社交关系，实现社群发现和活动推荐等更高维度的校园社交服务。")

    add_pb(doc)
    add_h(doc, "参考文献", 1)
    refs = [
        "[1] MacQueen J. Some methods for classification and analysis of multivariate observations[C]. Proceedings of the Fifth Berkeley Symposium on Mathematical Statistics and Probability, 1967, 1: 281-297.",
        "[2] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.",
        "[3] Chawla N V, Bowyer K W, Hall L O, et al. SMOTE: Synthetic minority over-sampling technique[J]. Journal of Artificial Intelligence Research, 2002, 16: 321-357.",
        "[4] Stone M. Cross-validatory choice and assessment of statistical predictions[J]. Journal of the Royal Statistical Society: Series B (Methodological), 1974, 36(2): 111-133.",
        "[5] Rousseeuw P J. Silhouettes: A graphical aid to the interpretation and validation of cluster analysis[J]. Journal of Computational and Applied Mathematics, 1987, 20: 53-65.",
        "[6] Myers I B, McCaulley M H, Quenk N L, et al. MBTI Manual: A Guide to the Development and Use of the Myers-Briggs Type Indicator[M]. 3rd ed. Palo Alto: Consulting Psychologists Press, 1998.",
        "[7] Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.",
        "[8] Patki N, Wedge R, Veeramachaneni K. The Synthetic Data Vault[C]. IEEE International Conference on Data Science and Advanced Analytics (DSAA), 2016: 399-410.",
        "[9] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.",
        "[10] 李航. 统计学习方法[M]. 2版. 北京: 清华大学出版社, 2019.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.add_run(ref).font.size = Pt(11)

    out = os.path.join(BASE, "机器学习课程报告_校园圈层聚类与推荐系统.docx")
    doc.save(out)
    print(f"DONE: {out}")

print("gen_course_report defined OK")
def gen_exp_report():
    doc = setup_doc()
    title = "基于K-Means与随机森林的\n校园搭子推荐系统实现与应用"
    doc = add_cover(doc, "机器学习实验", "2351303", title, is_exp=True)
    add_personnel(doc)
    add_pb(doc); add_abstract(doc, "experiment"); add_pb(doc)

    add_h(doc, "一、实验环境与技术准备", 1)
    add_p(doc, "开发环境：Windows 11 + Visual Studio Code + Anaconda（虚拟环境ML_Course），Python 3.x。核心依赖库：pandas（v1.5+，数据读取与DataFrame清洗操作）、numpy（v1.23+，数值计算与矩阵运算）、scikit-learn（v1.2+，提供K-Means、RandomForestClassifier、StandardScaler、PCA、LeaveOneOut交叉验证、cosine_similarity等核心API）、matplotlib（v3.7+）与seaborn（v0.12+，用于肘部法则图、雷达图、热力图、PCA散点图、混淆矩阵、ROC曲线、指标对比柱状图、特征重要性图共8类可视化输出，每张图均以dpi=150保存为PNG）、imbalanced-learn（v0.11+，提供SMOTENC混合型数据过采样）、streamlit（v1.25+，Web推荐系统前端框架）、flask + flask_cors（后端API与校园搭子推荐系统部署）、python-docx（报告自动化生成）、pickle（模型持久化存储）。")
    add_p(doc, "数据集来源：自行设计的MBTI性格与大学生消费行为关联问卷，通过问卷星平台在线发放与回收。问卷涵盖基础信息（性别、年级、生活费来源与区间）、细分消费结构（餐饮伙食、社交娱乐、服饰美妆、学习发展、数码游戏五维）、消费行为与认知（决策习惯、看重因素、兴趣投入意愿）、MBTI性格快速自测（E/I、S/N、T/F、J/P四个维度各1题）和社交偏好共5大模块19道题目。原始数据69份，经清洗后得62条有效记录，保存为data_cleaned.csv（UTF-8编码），后续所有实验均基于此数据集。")

    add_pb(doc)
    add_h(doc, "二、实验原理与实验方案", 1)
    add_p(doc, "本实验涉及两种核心算法的工程实现与四种关键技术的应用：（1）K-Means聚类——基于欧氏距离的划分式聚类，核心流程为StandardScaler标准化、K=2到8遍历、肘部法则（Inertia拐点）+轮廓系数、确定K=4、n_init=10执行最终聚类。输入为7维消费特征（5维消费明细+月预算+兴趣投入），输出为0-3的群体标签。（2）随机森林分类器——Bagging集成方法，n_estimators=100棵树投票，正则化参数max_depth=5（防止过拟合）、min_samples_leaf=3（保证叶节点代表性）、class_weight=balanced（自动加权缓解不均衡）。以12维特征（7维消费+3维人口学+2维消费行为）为输入，分别预测MBTI四个维度。关键技术依次为：LOOCV留一法——将62份样本逐一作为测试集，其余61份训练，共进行62轮训练-评估循环；SMOTENC——在每轮LOOCV训练前对训练集进行过采样，类别特征索引为[7,8,9,11]（decision_style、priority_factor、gender、income_source）；Bootstrap 95% CI——对最终预测结果进行2000次有放回重采样，计算2.5%和97.5%分位数作为置信区间上下界；Dummy Classifier——以most_frequent策略（始终预测训练集中最多的类别）建立基线，评估模型是否真正超越简单规则。")
    add_p(doc, "实验方案按三阶段组织：阶段一执行phase1_clustering.py，依次完成数据加载（pd.read_excel）、列名映射（COLUMN_MAP字典）、数据清洗（duration_sec<30s标记+五维消费全同标记，合并剔除）、StandardScaler标准化、肘部法则+轮廓系数图、K=4 K-Means聚类、群体画像（消费均值+MBTI分布）、4张可视化图（肘部+轮廓系数、雷达图、热力图、PCA散点图）、保存模型和清洗数据。阶段二执行phase2_mbti_classifier.py，依次完成数据加载、12维特征定义+类别特征索引、四维LOOCV循环（每维：LeaveOneOut.split、SMOTENC、StandardScaler、RF训练、预测、累积结果）、多指标计算（Acc、Macro F1、Precision、Recall、AUC-ROC）、Bootstrap CI、Dummy基线、4张可视化图（混淆矩阵、ROC曲线、指标对比、特征重要性）、关键发现输出、保存模型。阶段三执行phase3_app.py（Streamlit）或app.py（Flask），加载全部模型、用户输入消费画像、K-Means聚类获得群体标签、RF预测MBTI、Cosine相似度+MBTI互补+预算匹配+场景契合、输出TOP-5搭子。")

    add_pb(doc)
    add_h(doc, "三、实验过程与代码实现", 1)
    add_p(doc, "实验代码分为三个独立Python脚本和两个Web应用文件，总计约550行核心逻辑代码。所有代码均包含详细的中文注释，采用模块化设计原则，各阶段通过CSV和pickle文件进行数据交换，实现松耦合。")
    add_h(doc, "3.1 阶段一：数据清洗与聚类（phase1_clustering.py，约250行）", 2)
    add_p(doc, "首先使用pd.read_excel加载原始69条数据，通过COLUMN_MAP字典将问卷题目映射为英文变量名。接着解析填写时间字符串（通过str.replace去除秒字后astype(int)转换），标记两类脏数据：fast_mask（duration_sec<30秒，即乱填样本）和all_same_mask（5维消费答案完全相同，可能为应付型填写）。特别地，移除了MBTI四维全同的清洗规则——经审查发现全部11条被标记记录均为真实ISTJ类型（全选1），不应作为脏数据剔除。合并dirty_mask后使用~dirty_mask布尔索引获得干净数据集（62条）。聚类特征选取consume_food、consume_social、consume_fashion、consume_study、consume_digital加monthly_budget加hobby_spend共7维，经StandardScaler标准化后以K=2到8遍历K-Means（n_init=10, random_state=42），记录每次的inertia和silhouette_score。最终K=4执行聚类，计算cluster_profile（群体均值），使用LABEL_MAP字典将数值映射为可读标签，并基于均值的启发式规则为各群体命名。所有可视化使用matplotlib/seaborn，以dpi=150保存高分辨率PNG图片。")
    add_h(doc, "3.2 阶段二：MBTI分类（phase2_mbti_classifier.py，约350行）", 2)
    add_p(doc, "特征空间扩展至12维，定义CATEGORICAL_INDICES = [7,8,9,11]指定混合型数据中的类别特征位置。对四个MBTI维度分别执行完整的LOOCV流程：使用LeaveOneOut().split(X_all)生成62组(train, test)索引对，每轮训练中根据y_train的类别分布动态计算smote_k_neighbors = min(3, max(1, min_class_count-1))——当某类只剩1个样本时自动跳过SMOTE。SMOTENC的categorical_features参数精确传入类别特征索引列表。RF使用固定超参数（n_estimators=100, max_depth=5, min_samples_leaf=3, class_weight=balanced, random_state=42）以保证可复现性。评估指标的计算分为两步：先在LOOCV累积的y_true_all和y_pred_all上计算点估计（Accuracy、Macro F1、Precision、Recall），再通过bootstrap_ci函数进行2000次Bootstrap重采样获取95%置信区间。AUC的计算根据维度分类数自动选择：二分类直接用roc_auc_score，三分类（E/I维度）使用label_binarize + macro averaging。")
    add_h(doc, "3.3 阶段三：推荐系统（app.py/phase3_app.py，约250行）", 2)
    add_p(doc, "Flask后端（app.py）使用@app.route装饰器定义RESTful API接口，/api/recommend接收POST请求，解析JSON中的12维用户消费特征，先后调用K-Means（predict）、RF四维（predict_proba取概率最大类别）、计算Cosine相似度矩阵、综合MBTI互补得分（E/I互补充1，同类型扣分；T/F互补；S/N和J/P相似加分）+预算匹配得分（差值小于等于1档得1分）+场景契合得分（匹配3分，部分1分）、加权求和、返回TOP-5推荐结果。Streamlit前端（phase3_app.py）使用@st.cache_resource和@st.cache_data装饰器缓存模型和数据，通过st.slider（连续特征）、st.selectbox（类别特征）构建输入表单，以plotly生成交互式雷达图，推荐结果以卡片形式展示，包含匹配得分、消费群体标签、MBTI类型和推荐理由。")

    add_pb(doc)
    add_h(doc, "四、实验结果与可视化分析", 1)
    add_h(doc, "4.1 聚类结果", 2)
    add_p(doc, "肘部法则图（output_elbow_silhouette.png）显示K=2到3到4时Inertia下降速率从33.0到21.5到14.7，K=4处出现明显拐点，此后速率趋于平缓（K=5到6仅下降约6个单位）。轮廓系数在K=4时达到局部峰值0.1775（高于K=3的0.1655和K=5的0.1708），两指标联合判定最优K=4。K=4聚类结果显示：群体0（高消费数码型，n=12）预算最高（4.08），数码消费突出（3.17），兴趣投入高（2.08），主流MBTI为XNTP。群体1（节俭生存型，n=23）五维消费全部垫底，月预算均值仅2.13，主流MBTI为ISFP。群体2（社交体验型，n=14）社交消费最高（2.64），服饰消费第二（2.21），兴趣投入最低（1.14），主流MBTI为XSFP。群体3（自我提升型，n=13）学习消费最高（3.0），服饰消费同样最高（2.77），兴趣投入最高（2.38），主流MBTI为XNFP。")
    add_p(doc, "可视化成果：雷达图（output_radar.png）以7维特征为轴，四条不同颜色的折线直观展示了四类群体在消费结构上的鲜明差异——群体0在数码轴突出、群体1全面收缩、群体2在社交轴隆起、群体3在学习和兴趣轴双双突出。热力图（output_heatmap.png）以颜色深浅呈现群体与特征的均值矩阵，YlOrRd配色方案使得高消费区域一目了然。PCA散点图（output_pca.png）将7维特征投影至2维平面（PC1+PC2解释约55%方差），四类群体用不同颜色区分，黑色X标记聚类中心，验证了群体间存在适度分离但边界不完全清晰的特点。")
    add_h(doc, "4.2 MBTI分类结果", 2)
    add_p(doc, "四维分类的完整评估结果汇总如下：J/P生活风格维度表现最佳（Acc=58.06%, F1=0.580, AUC=0.647, 超Dummy基线+6.45%），两类F1均衡（J=0.57, P=0.59），选购看重因素为最强特征（重要性0.167）。S/N信息处理维度Acc最高（61.29%），但Macro F1仅0.495且AUC偏低（0.437），低于Dummy基线6.45%，社交娱乐消费为最强特征（重要性0.218）。E/I精力恢复维度受9份外倾样本限制，E类完全无法预测（F1=0.00），整体Acc仅48.39%，数码消费和性别分别为重要性第2和第1特征。T/F决策标准维度Acc最低（45.16%），低于Dummy基线11.29%，服饰美妆消费为最强特征（重要性0.191），支持情感型更关注外在形象的假设。四维平均Acc为53.23%，平均AUC为0.470。")
    add_p(doc, "可视化成果：混淆矩阵（output_confusion_matrices.png）以4个子图展示每维分类的预测分布，对角线颜色深浅直观反映分类准确情况——J/P维度对角线颜色最深，E/I维度E类对角线几乎无色。ROC曲线（output_roc_curves.png）绘制了三个二分类维度（S/N、T/F、J/P）的ROC曲线，J/P的曲线最接近左上角（AUC=0.647），T/F的曲线徘徊在对角线附近（AUC=0.505）。指标对比柱状图（output_metrics_comparison.png）将Acc、F1、AUC三个指标与Dummy基线并排展示，清晰呈现了各维度的相对优势和不足。特征重要性条形图（output_feature_importance.png）以4个子图展示12维特征对每个MBTI维度的贡献，社交娱乐、性别、服饰美妆、看重因素分别在S/N、E/I、T/F、J/P维度中排名第一。")
    add_h(doc, "4.3 推荐系统演示", 2)
    add_p(doc, "Streamlit Web应用实现了端到端的用户交互流程：用户通过滑块和下拉框输入12维消费画像后，系统实时返回聚类群体标签（如社交体验型）、MBTI四维预测结果（如ESFP）和TOP-5搭子推荐（含匹配得分、群体标签、MBTI类型和推荐理由）。在62份样本的内部测试中，推荐结果展现出合理的匹配逻辑：预算相近、消费结构相似的搭子排名靠前；MBTI互补的搭子（如E型推荐I型）获得额外加分；场景偏好匹配的用户排在同类场景用户前面。系统响应时间在Streamlit框架下保持在1秒以内（模型预加载+缓存机制），满足实时交互的体验要求。")

    add_pb(doc)
    add_h(doc, "五、实验总结与反思", 1)
    add_p(doc, "本实验成功完成以下目标：（1）掌握了K-Means聚类算法的完整实现流程，包括数据标准化、最优K值确定（肘部法则+轮廓系数联合验证）、聚类结果的多维可视化解读（雷达图、热力图、PCA散点图）。（2）深入实践了Random Forest分类器的训练、评估与调优，深刻理解了LOOCV在小样本场景下相较于传统train_test_split的显著优势。（3）掌握了SMOTENC对混合型数据类别不平衡的处理机制，以及Bootstrap置信区间对模型评估不确定性的量化方法。（4）运用Streamlit和Flask双框架将训练好的模型部署为可交互的Web推荐系统，实现了从数据到产品的完整机器学习项目闭环。（5）通过Dummy Classifier基线的引入，培养了批判性评估模型性能的科学态度——避免被高Accuracy误导。")
    add_p(doc, "反思实验过程，主要局限性与改进方向包括：（1）数据层面——62份样本在机器学习中属于极小样本，E/I维度仅9份E类样本是最大的瓶颈。未来可尝试使用SDV GaussianCopulaSynthesizer生成合成数据扩充至200-300条，或引入校园一卡通消费流水等客观数据源。（2）特征工程——当前的12维特征完全基于问卷设计，缺乏交叉特征和组合特征。可尝试多项式特征生成（PolynomialFeatures）或基于聚类的特征构造（Cluster Distance Features）来增强模型表达能力。（3）模型选择——当前仅使用RF一种分类器，未进行模型对比。后续可引入XGBoost、LightGBM、SVM、MLP等多种模型，通过统计检验（如McNemar检验）判断差异是否显著。（4）超参数调优——当前RF参数（max_depth=5等）由经验设定，可通过GridSearchCV或Optuna贝叶斯优化在LOOCV框架内进行超参数搜索。（5）聚类方法——K-Means对簇形状的球形假设在实际数据中可能不完全满足，可尝试DBSCAN（发现任意形状的簇）或GMM（支持非等方差的高斯混合模型）进行对比实验。")
    add_p(doc, "在技能成长方面，本次实验使团队成员在实践中掌握了：Python机器学习生态（scikit-learn + pandas + matplotlib）的完整使用、小样本场景下的模型评估策略设计、数据清洗的工程化方法（列名映射+脏数据标记+布尔索引过滤）、模型持久化和Web部署（pickle + Flask/Streamlit）等核心技能，为后续从事机器学习相关研究和工程开发奠定了坚实基础。")

    out = os.path.join(BASE, "机器学习实验课程报告_校园圈层聚类与推荐系统.docx")
    doc.save(out)
    print(f"DONE: {out}")

print("gen_exp_report defined OK")
if __name__ == "__main__":
    print("=" * 60)
    print("Generating Course Report...")
    gen_course_report()
    print("=" * 60)
    print("Generating Experiment Report...")
    gen_exp_report()
    print("=" * 60)
    print("ALL DONE! Both reports generated successfully.")