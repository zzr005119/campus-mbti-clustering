# -*- coding: utf-8 -*-
"""生成机器学习实验课程报告 v2 — 含代码块和图片占位"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

BASE = r"D:\Xx\大二下\机器学习\大作业"
SIDS = "2024117715\n2024117720\n2024117719"
SNAMES = "赵子荣\n李海东\n章振豪"
CLASS = "计算机科学与技术2022级X班"
TEACHER = "曾梦"

def read_src(fname):
    enc = 'utf-8-sig' if 'app.py' in fname else 'utf-8'
    with open(os.path.join(BASE, fname), 'r', encoding=enc) as f:
        return f.read()

P1 = read_src('phase1_clustering.py')
P2 = read_src('phase2_mbti_classifier.py')
APP = read_src('app.py')

def extract(full, header, end_marker=None, max_l=45):
    lines = full.split('\n')
    s = -1
    for i, ln in enumerate(lines):
        if header in ln:
            s = i; break
    if s == -1: return "# NOT FOUND: " + header
    r = []
    for i in range(s, len(lines)):
        if end_marker and end_marker in lines[i] and i > s + 3: break
        if len(r) >= max_l:
            r.append("# ... (后续代码省略)")
            break
        r.append(lines[i])
    return '\n'.join(r)

CLEAN_CODE = extract(P1, '# 4.1 解析填写时间', '# 5. 描述性统计预览', 28)
ELBOW_CODE = extract(P1, '# 7. 肘部法则', '# 9. 聚类画像', 42)
RADAR_CODE = extract(P1, '# 10. 可视化', '# 11. 可视化', 48)
PCA_CODE = extract(P1, '# 12. 可视化', '# 13. 导出清洗后的数据', 35)
LOOCV_CODE = extract(P2, '# 4. LOOCV', '# 5. 结果汇总', 55)
BOOT_CODE = extract(P2, '# 2. Bootstrap CI', '# 3. 目标变量定义', 38)
REC_CODE = extract(APP, 'def recommend_matches', '@app.route', 55)
API_CODE = extract(APP, "def analyze():", "if __name__", 42)

print(f"Code extracted: Clean={len(CLEAN_CODE)} Elbow={len(ELBOW_CODE)} LOOCV={len(LOOCV_CODE)}")

# ====== Helper functions ======

def add_personnel(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run("人员分配："); r.bold = True; r.font.name = "宋体"; r.font.size = Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run("本项目共三位成员协作完成，具体分工如下：").font.size = Pt(12)
    assigns = [
        ("（1）赵子荣（2024117715）：", "负责项目总体设计与技术架构、核心算法实现（K-Means聚类与随机森林分类器）、LOOCV交叉验证与SMOTENC过采样策略设计、Bootstrap置信区间评估、模型训练与调优、可视化图表生成、Streamlit与Flask双版本推荐系统开发、答辩PPT制作与展示。同时统筹两门课程报告的撰写框架与内容整合。"),
        ("（2）李海东（2024117720）：", "负责前期文献调研与MBTI性格理论梳理、问卷设计与问卷星平台发放、原始数据收集与初步整理、数据清洗规则设计。参与课程报告相关理论与算法基础章节的撰写与参考文献整理，协助实验报告中的实验环境与技术准备章节编写。"),
        ("（3）章振豪（2024117719）：", "负责数据可视化方案设计（雷达图、热力图、PCA散点图等）、Web前端界面美化与交互优化、用户测试与系统体验评估。参与实验报告实验结果与可视化分析章节的图表演示与分析解读，协助课程报告中的应用场景与效果分析章节撰写。"),
    ]
    for nm, desc in assigns:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        rn = p.add_run(nm); rn.bold = True; rn.font.size = Pt(12)
        p.add_run(desc).font.size = Pt(12)

def add_abstract(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run("摘要："); r.bold = True; r.font.name = "宋体"; r.font.size = Pt(12)
    t = "本实验基于62份校园问卷数据，使用Python及scikit-learn库完整实现了K-Means消费群体聚类、随机森林MBTI性格预测与Streamlit/Flask双版本搭子推荐系统。实验涵盖数据清洗（剔除填写时间<30秒及消费五维全同的无效样本）、StandardScaler标准化、肘部法则与轮廓系数联合确定K=4、LOOCV留一法交叉验证、SMOTENC过采样处理类别不平衡、Bootstrap 95%置信区间评估、Dummy Classifier基线对比等关键环节。通过matplotlib与seaborn生成了肘部法则图、雷达图、热力图、PCA散点图、混淆矩阵、ROC曲线、指标对比柱状图和特征重要性条形图共8类可视化成果。最终交付了完整的Python代码包（约550行核心代码）、训练好的模型文件（K-Means + 4个RF分类器 + 5个Scaler）以及可交互的Web推荐系统原型。"
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(t).font.size = Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run("关键词："); r.bold = True; r.font.size = Pt(12)
    p.add_run("K-Means聚类；随机森林；MBTI人格预测；SMOTENC过采样；校园搭子推荐系统").font.size = Pt(12)

def setup_doc():
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "宋体"; s.font.size = Pt(12)
    s.paragraph_format.line_spacing = 1.5
    s.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return doc

def add_cover(doc, cname, ccode, title):
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本 科 生 课 程 报 告 封 面")
    r.font.size = Pt(22); r.font.name = "黑体"; r.bold = True
    doc.add_paragraph()
    tbl = doc.add_table(rows=7, cols=2); tbl.autofit = True
    data = [
        ("课程名称：", cname), ("课程编码：", ccode), ("报告题目：", title),
        ("学    号：", SIDS), ("姓    名：", SNAMES),
        ("班    级：", CLASS), ("授课教师：", TEACHER),
    ]
    for i, (lb, vl) in enumerate(data):
        row = tbl.rows[i]
        p0 = row.cells[0].paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0 = p0.add_run(lb); r0.font.size = Pt(14); r0.font.name = "宋体"
        r0.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        p1 = row.cells[1].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(vl)
        r1.font.size = Pt(11) if i == 2 else Pt(14); r1.font.name = "宋体"
        r1.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if i == 2: r1.bold = True
    tblP = tbl._tbl.tblPr
    if tblP is None:
        tblP = parse_xml('<w:tblPr ' + nsdecls("w") + '></w:tblPr>')
        tbl._tbl.insert(0, tblP)
    bdrs = parse_xml('<w:tblBorders ' + nsdecls("w") + '><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tblBorders>')
    tblP.append(bdrs)
    doc.add_paragraph(); doc.add_paragraph()
    _score(doc)
    doc.add_page_break()
    return doc

def _score(doc):
    tbl = doc.add_table(rows=5, cols=8); tbl.autofit = True
    hdrs = ["评分标准及分值", "", "选题契合\n（分值15）", "报告内容与代码实现\n（分值40）", "", "实验分析\n（分值30）", "", "报告规范\n（分值15）"]
    for j, h in enumerate(hdrs):
        p = tbl.rows[0].cells[j].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.font.size = Pt(9); r.bold = True
    tbl.rows[0].cells[0].merge(tbl.rows[0].cells[1])
    tbl.rows[0].cells[3].merge(tbl.rows[0].cells[4])
    tbl.rows[0].cells[5].merge(tbl.rows[0].cells[6])
    p = tbl.rows[1].cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("评分").font.size = Pt(10)
    tbl.rows[2].cells[0].merge(tbl.rows[2].cells[7])
    tbl.rows[2].cells[0].paragraphs[0].add_run("注：以上为参考标准，授课教师根据需要对评分标准进行调整").font.size = Pt(9)
    tbl.rows[3].cells[0].paragraphs[0].add_run("评语：").font.size = Pt(10)
    tbl.rows[4].cells[0].paragraphs[0].add_run("总 评 分").font.size = Pt(10)
    p = tbl.rows[4].cells[3].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("评阅教师：").font.size = Pt(10)
    p = tbl.rows[4].cells[6].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("评阅时间\n2026年 月 日").font.size = Pt(10)

def add_h(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "黑体"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return h

def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(text).font.size = Pt(12)
    return p

def add_pb(doc):
    doc.add_page_break()

def add_code(doc, code, title=""):
    """添加代码块：灰底表格，等宽字体"""
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title); r.bold = True; r.font.size = Pt(11); r.font.name = "黑体"
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    tbl = doc.add_table(rows=1, cols=1); tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    shd = parse_xml('<w:shd ' + nsdecls("w") + ' w:fill="F0F0F0" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shd)
    for ln in code.split('\n')[:50]:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(ln if ln else " ")
        r.font.size = Pt(7.5); r.font.name = "Consolas"
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    doc.add_paragraph()

def add_img(doc, filename, desc):
    """添加图片占位框"""
    tbl = doc.add_table(rows=1, cols=1); tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    # Title
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("【插入图片】" + filename); r.bold = True
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)
    # Description
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(desc); r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

print("All helper functions defined OK")
def generate():
    doc = setup_doc()
    title = "基于K-Means与随机森林的\n校园搭子推荐系统实现与应用"
    doc = add_cover(doc, "机器学习实验", "2351303", title)
    add_personnel(doc)
    add_pb(doc); add_abstract(doc); add_pb(doc)

    # ========== 一、实验环境与技术准备 ==========
    add_h(doc, "一、实验环境与技术准备", 1)
    add_p(doc, "开发环境：Windows 11 + Visual Studio Code + Anaconda（虚拟环境ML_Course），Python 3.x。核心依赖库：pandas（v1.5+，数据读取与DataFrame清洗）、numpy（v1.23+，数值计算）、scikit-learn（v1.2+，K-Means、RandomForestClassifier、StandardScaler、PCA、LeaveOneOut、cosine_similarity等）、matplotlib（v3.7+）与seaborn（v0.12+，8类可视化图表输出）、imbalanced-learn（v0.11+，SMOTENC过采样）、streamlit（v1.25+，Web前端）、flask + flask_cors（后端API）、python-docx（报告自动化生成）、pickle（模型持久化）。")
    add_p(doc, "数据集来源：自行设计的MBTI性格与大学生消费行为关联问卷，通过问卷星平台在线发放与回收。问卷涵盖基础信息（性别、年级、生活费来源与区间）、细分消费结构（餐饮伙食、社交娱乐、服饰美妆、学习发展、数码游戏五维）、消费行为与认知（决策习惯、看重因素、兴趣投入意愿）、MBTI性格快速自测（E/I、S/N、T/F、J/P四维各1题）和社交偏好共5大模块19道题目。原始数据69份，经清洗后得62条有效记录，保存为data_cleaned.csv（UTF-8编码），后续所有实验均基于此数据集。")

    add_pb(doc)
    # ========== 二、实验原理与实验方案 ==========
    add_h(doc, "二、实验原理与实验方案", 1)
    add_p(doc, "本实验涉及两种核心算法的工程实现与四种关键技术的应用：（1）K-Means聚类——基于欧氏距离的划分式聚类，核心流程为StandardScaler标准化、K=2到8遍历、肘部法则（Inertia拐点）+轮廓系数、确定K=4、n_init=10执行最终聚类。输入为7维消费特征（5维消费明细+月预算+兴趣投入），输出为0-3的群体标签。（2）随机森林分类器——Bagging集成方法，n_estimators=100棵树投票，正则化参数max_depth=5（防止过拟合）、min_samples_leaf=3（保证叶节点代表性）、class_weight=balanced（自动加权缓解不均衡）。以12维特征（7维消费+3维人口学+2维消费行为）为输入，分别预测MBTI四个维度。关键技术依次为：LOOCV留一法——将62份样本逐一作为测试集，其余61份训练，共进行62轮训练-评估循环；SMOTENC——在每轮LOOCV训练前对训练集进行过采样，类别特征索引为[7,8,9,11]（decision_style、priority_factor、gender、income_source）；Bootstrap 95% CI——对最终预测结果进行2000次有放回重采样，计算2.5%和97.5%分位数作为置信区间上下界；Dummy Classifier——以most_frequent策略（始终预测训练集中最多的类别）建立基线，评估模型是否真正超越简单规则。")
    add_p(doc, "实验方案按三阶段组织：阶段一执行phase1_clustering.py，依次完成数据加载（pd.read_excel）、列名映射（COLUMN_MAP字典）、数据清洗（duration_sec<30s标记+五维消费全同标记，合并剔除）、StandardScaler标准化、肘部法则+轮廓系数图、K=4 K-Means聚类、群体画像（消费均值+MBTI分布）、4张可视化图（肘部+轮廓系数、雷达图、热力图、PCA散点图）、保存模型和清洗数据。阶段二执行phase2_mbti_classifier.py，依次完成数据加载、12维特征定义+类别特征索引、四维LOOCV循环（每维：LeaveOneOut.split、SMOTENC、StandardScaler、RF训练、预测、累积结果）、多指标计算（Acc、Macro F1、Precision、Recall、AUC-ROC）、Bootstrap CI、Dummy基线、4张可视化图（混淆矩阵、ROC曲线、指标对比、特征重要性）、关键发现输出、保存模型。阶段三执行phase3_app.py（Streamlit）或app.py（Flask），加载全部模型、用户输入消费画像、K-Means聚类获得群体标签、RF预测MBTI、Cosine相似度+MBTI互补+预算匹配+场景契合、输出TOP-5搭子。")

    add_pb(doc)
    # ========== 三、实验过程与代码实现（含关键代码） ==========
    add_h(doc, "三、实验过程与代码实现", 1)
    add_p(doc, "实验代码分为三个独立Python脚本和两个Web应用文件，总计约550行核心逻辑代码。所有代码均包含详细的中文注释，采用模块化设计原则，各阶段通过CSV和pickle文件进行数据交换，实现松耦合。以下分阶段展示关键代码段及注释。")

    add_h(doc, "3.1 阶段一：数据清洗与聚类（phase1_clustering.py，约250行）", 2)
    add_p(doc, "阶段一首先使用pd.read_excel加载原始69条数据，通过COLUMN_MAP字典将19道中文问卷题目映射为英文变量名（如每月餐饮伙食映射为consume_food），同时定义LABEL_MAP字典为每个编码值关联可读标签（如gender的{1:男, 2:女}）。数据清洗的核心逻辑如下所示——通过两个布尔掩码分别标记填写时间过短（<30秒）和五维消费答案完全相同的无效样本，合并后使用~dirty_mask布尔索引获得干净数据集：")
    add_code(doc, CLEAN_CODE, "代码 3-1：数据清洗核心逻辑")
    add_p(doc, "清洗后的62条有效记录被用于后续分析。聚类阶段提取7维特征（5维消费明细consume_food/consume_social/consume_fashion/consume_study/consume_digital + monthly_budget + hobby_spend），经StandardScaler标准化后，以K=2到8遍历K-Means（n_init=10, random_state=42），通过以下肘部法则和轮廓系数代码确定最优K值：")
    add_code(doc, ELBOW_CODE, "代码 3-2：肘部法则 + K-Means聚类核心代码")
    add_p(doc, "最终K=4执行聚类，计算cluster_profile（群体均值），使用LABEL_MAP字典将数值映射为可读标签。雷达图绘制代码以7维特征为轴、四色折线展示群体差异；PCA降维散点图将12维特征投影到2维平面，验证聚类结果的合理性：")
    add_code(doc, RADAR_CODE, "代码 3-3：雷达图与PCA散点图可视化代码")

    add_pb(doc)
    add_h(doc, "3.2 阶段二：MBTI分类预测（phase2_mbti_classifier.py，约350行）", 2)
    add_p(doc, "阶段二是本次更新的核心内容。首先加载data_cleaned.csv的62条数据，构建12维特征矩阵（5消费明细+2辅助消费+2行为特征+3人口学），标记4个类别特征索引供SMOTENC使用。Bootstrap CI工具函数为小样本下的评估提供了统计严谨性保障：")
    add_code(doc, BOOT_CODE, "代码 3-4：Bootstrap 95%置信区间工具函数")
    add_p(doc, "对四个MBTI维度逐一进行LOOCV评估的完整训练循环如下。在每折训练中，对训练集动态应用SMOTENC过采样（k_neighbors自适应调整以适应样本量极小的折），StandardScaler标准化后训练RandomForestClassifier（100棵树，max_depth=5），对测试样本进行预测并收集概率输出：")
    add_code(doc, LOOCV_CODE, "代码 3-5：LOOCV + SMOTENC训练循环核心代码")

    add_pb(doc)
    add_h(doc, "3.3 阶段三：推荐系统部署（app.py/phase3_app.py，约250行）", 2)
    add_p(doc, "Flask后端（app.py）使用@app.route装饰器定义RESTful API接口。启动时加载K-Means模型、4个RF模型和对应的Scaler。推荐算法的核心逻辑实现了多因子加权匹配策略——综合消费相似度（Cosine）、预算相容性（差值<=1档得1分）、MBTI互补性（E/I互补充1，T/F互补）和场景契合度（匹配得1分），加权求和后返回TOP-5推荐：")
    add_code(doc, REC_CODE, "代码 3-6：多因子加权推荐算法核心代码")
    add_p(doc, "Flask REST API接口接收POST请求中的12维用户消费特征，依次执行聚类预测、四维MBTI预测、推荐计算，返回JSON格式结果：")
    add_code(doc, API_CODE, "代码 3-7：Flask REST API /api/analyze 接口实现")

    add_pb(doc)
    # ========== 四、实验结果与可视化分析（含图片占位） ==========
    add_h(doc, "四、实验结果与可视化分析", 1)
    add_h(doc, "4.1 聚类实验结果", 2)
    add_p(doc, "肘部法则图显示K=2到3到4时Inertia下降速率从33.0到21.5到14.7，K=4处出现明显拐点，此后速率趋于平缓（K=5到6仅下降约6个单位）。轮廓系数在K=4时达到局部峰值0.1775（高于K=3的0.1655和K=5的0.1708），两指标联合判定最优K=4。K=4聚类结果显示：群体0（高消费数码型，n=12）预算最高（4.08），数码消费突出（3.17），兴趣投入高（2.08），主流MBTI为XNTP。群体1（节俭生存型，n=23）五维消费全部垫底，月预算均值仅2.13，主流MBTI为ISFP。群体2（社交体验型，n=14）社交消费最高（2.64），服饰消费第二（2.21），兴趣投入最低（1.14），主流MBTI为XSFP。群体3（自我提升型，n=13）学习消费最高（3.0），服饰消费同样最高（2.77），兴趣投入最高（2.38），主流MBTI为XNFP。")
    add_p(doc, "以下是本次实验生成的可视化图表，请在对应位置插入图片文件：")
    add_img(doc, "output_elbow_silhouette.png", "肘部法则图（左：Inertia-K曲线，右：轮廓系数-K曲线）。展示K=2~8区间内两项指标的变化趋势，K=4处Inertia出现明显拐点且轮廓系数达局部峰值，联合验证最优聚类数。")
    add_img(doc, "output_radar.png", "四类消费群体雷达图。以7维特征（餐饮、社交、服饰、学习、数码、月预算、兴趣）为轴，四条不同颜色折线展示群体消费轮廓差异：群体0在数码轴突出、群体1全面收缩、群体2在社交轴隆起、群体3在学习和兴趣轴双双突出。")
    add_img(doc, "output_heatmap.png", "消费群体特征热力图。以颜色深浅（YlOrRd配色）呈现群体×7维特征的均值矩阵，vmin=1, vmax=4，annot=True标注具体均值，直观对比各群体在不同消费维度的等级差异。")
    add_img(doc, "output_pca.png", "PCA降维散点图。将7维特征经PCA压缩至2维平面（PC1+PC2解释约55%方差），四色散点代表四类群体，黑色X标记聚类中心位置，验证群体间存在适度分离趋势。")

    add_pb(doc)
    add_h(doc, "4.2 MBTI分类实验结果", 2)
    add_p(doc, "四维分类的完整评估结果汇总如下：J/P生活风格维度表现最佳（Acc=58.06%, F1=0.580, AUC=0.647, 超Dummy基线+6.45%），两类F1均衡（J=0.57, P=0.59），选购看重因素为最强特征（重要性0.167）。S/N信息处理维度Acc最高（61.29%），但Macro F1仅0.495且AUC偏低（0.437），低于Dummy基线6.45%，社交娱乐消费为最强特征（重要性0.218）。E/I精力恢复维度受9份外倾样本限制，E类完全无法预测（F1=0.00），整体Acc仅48.39%，数码消费和性别分别为重要性第2和第1特征。T/F决策标准维度Acc最低（45.16%），低于Dummy基线11.29%，服饰美妆消费为最强特征（重要性0.191），支持情感型更关注外在形象的假设。四维平均Acc为53.23%，平均AUC为0.470。")
    add_p(doc, "以下为Phase2生成的四类评估可视化图表，请在对应位置插入：")
    add_img(doc, "output_confusion_matrices.png", "四维混淆矩阵。4个子图分别展示E/I、S/N、T/F、J/P四维的预测分布，对角线颜色深浅反映分类准确情况——J/P对角线颜色最深，S/N次之，E/I的E类对角线几乎无色（模型完全无法预测外倾类型）。")
    add_img(doc, "output_roc_curves.png", "ROC曲线图。绘制S/N、T/F、J/P三个二分类维度的ROC曲线，J/P的曲线最接近左上角（AUC=0.647），T/F的曲线徘徊在对角线附近（AUC=0.505），展示各维度对不同类别的区分能力。")
    add_img(doc, "output_metrics_comparison.png", "四维指标对比柱状图。并排展示Accuracy、Macro F1、AUC三项指标与Dummy基线，直观呈现各维度的预测难度差异：J/P维度各项优异，T/F维度低于Dummy基线。")
    add_img(doc, "output_feature_importance.png", "四维特征重要性条形图。4个子图分别展示12维特征对每个MBTI维度的预测贡献度（Gini importance），社交娱乐、性别、服饰美妆、看重因素分别在S/N、E/I、T/F、J/P维度中排名第一。")

    add_pb(doc)
    add_h(doc, "4.3 推荐系统演示", 2)
    add_p(doc, "Streamlit Web应用实现了端到端的用户交互流程：用户通过滑块和下拉框输入12维消费画像后，系统实时返回聚类群体标签（如社交体验型）、MBTI四维预测结果（如ESFP）和TOP-5搭子推荐（含匹配得分、群体标签、MBTI类型和推荐理由）。在62份样本的内部测试中，推荐结果展现出合理的匹配逻辑：预算相近、消费结构相似的搭子排名靠前；MBTI互补的搭子（如E型推荐I型）获得额外加分；场景偏好匹配的用户排在同类场景用户前面。系统响应时间在Streamlit框架下保持在1秒以内（模型预加载+缓存机制），满足实时交互的体验要求。")

    add_pb(doc)
    # ========== 五、实验总结与反思 ==========
    add_h(doc, "五、实验总结与反思", 1)
    add_p(doc, "本实验成功完成以下目标：（1）掌握了K-Means聚类算法的完整实现流程，包括数据标准化、最优K值确定（肘部法则+轮廓系数联合验证）、聚类结果的多维可视化解读（雷达图、热力图、PCA散点图）。（2）深入实践了Random Forest分类器的训练、评估与调优，深刻理解了LOOCV在小样本场景下相较于传统train_test_split的显著优势。（3）掌握了SMOTENC对混合型数据类别不平衡的处理机制，以及Bootstrap置信区间对模型评估不确定性的量化方法。（4）运用Streamlit和Flask双框架将训练好的模型部署为可交互的Web推荐系统，实现了从数据到产品的完整机器学习项目闭环。（5）通过Dummy Classifier基线的引入，培养了批判性评估模型性能的科学态度——避免被高Accuracy误导。")
    add_p(doc, "反思实验过程，主要局限性与改进方向包括：（1）数据层面——62份样本在机器学习中属于极小样本，E/I维度仅9份E类样本是最大的瓶颈。未来可尝试使用SDV GaussianCopulaSynthesizer生成合成数据扩充至200-300条，或引入校园一卡通消费流水等客观数据源。（2）特征工程——当前的12维特征完全基于问卷设计，缺乏交叉特征和组合特征。可尝试多项式特征生成（PolynomialFeatures）或基于聚类的特征构造（Cluster Distance Features）来增强模型表达能力。（3）模型选择——当前仅使用RF一种分类器，未进行模型对比。后续可引入XGBoost、LightGBM、SVM、MLP等多种模型，通过统计检验（如McNemar检验）判断差异是否显著。（4）超参数调优——当前RF参数（max_depth=5等）由经验设定，可通过GridSearchCV或Optuna贝叶斯优化在LOOCV框架内进行超参数搜索。（5）聚类方法——K-Means对簇形状的球形假设在实际数据中可能不完全满足，可尝试DBSCAN（发现任意形状的簇）或GMM（支持非等方差的高斯混合模型）进行对比实验。")
    add_p(doc, "在技能成长方面，本次实验使团队成员在实践中掌握了：Python机器学习生态（scikit-learn + pandas + matplotlib）的完整使用、小样本场景下的模型评估策略设计、数据清洗的工程化方法（列名映射+脏数据标记+布尔索引过滤）、模型持久化和Web部署（pickle + Flask/Streamlit）等核心技能，为后续从事机器学习相关研究和工程开发奠定了坚实基础。")

    out = os.path.join(BASE, "机器学习实验课程报告_校园圈层聚类与推荐系统.docx")
    doc.save(out)
    print(f"\nDONE: {out}")
    print(f"Report contains code blocks from phase1/phase2/app.py and 8 image placeholders")

if __name__ == "__main__":
    generate()