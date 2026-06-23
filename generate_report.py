"""
项目报告生成脚本
运行前确保已安装 python-docx：conda activate ML_Course && pip install python-docx
运行：python generate_report.py
输出：项目报告_校园圈层聚类与推荐系统.docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# 标题样式
for i in range(1, 4):
    heading = doc.styles[f'Heading {i}']
    heading.font.name = '黑体'
    heading.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ========== 封面 ==========
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('基于MBTI心理特质与微观消费特征的\n校园圈层聚类与推荐系统')
run.font.name = '黑体'
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('机器学习课程大作业 · 项目报告')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(datetime.date.today().strftime('%Y年%m月'))
run.font.size = Pt(14)

doc.add_page_break()

# ========== 目录占位 ==========
doc.add_heading('目  录', level=1)
doc.add_paragraph('[此处生成自动目录 — Word 中按 Ctrl+A 后 F9 更新域]')
doc.add_page_break()

# ========== 第一章：项目概述 ==========
doc.add_heading('一、项目概述', level=1)

doc.add_heading('1.1 项目背景', level=2)
doc.add_paragraph(
    '在大学校园中，消费行为和性格特质是影响学生社交匹配的两个核心维度。'
    '一方面，消费水平与消费结构直接影响学生能否"玩到一起"——预算差异过大的搭子在聚餐、'
    '出游等场景中容易产生摩擦；另一方面，MBTI人格类型被广泛用于解释个体的社交偏好和决策模式，'
    'E/I维度决定了一个人是通过社交获取能量还是独处恢复精力，J/P维度则影响出行计划的灵活性。'
)
doc.add_paragraph(
    '本项目将无监督聚类、有监督分类和推荐系统相结合，设计并实现了一个"校园搭子推荐系统"。'
    '系统基于62份真实大学生问卷数据，通过K-Means算法将消费群体划分为四类，'
    '利用随机森林模型实现"看消费习惯预测MBTI性格"，最终综合消费相似度、MBTI互补性和场景偏好'
    '给出个性化的搭子匹配推荐。'
)

doc.add_heading('1.2 项目目标', level=2)
doc.add_paragraph('本项目的三个阶段性目标如下：')
doc.add_paragraph('（1）阶段一：对问卷数据进行清洗与标准化，使用K-Means无监督聚类算法识别校园消费群体类型，并通过热力图、雷达图等可视化手段展示群体特征。', style='List Bullet')
doc.add_paragraph('（2）阶段二：以消费明细和消费行为作为特征，以MBTI四个维度作为预测目标，训练随机森林分类器，采用留一法交叉验证（LOOCV）评估模型性能，并通过特征重要性分析验证"社交花销→E倾向"等核心假设。', style='List Bullet')
doc.add_paragraph('（3）阶段三：基于阶段一和阶段二的模型成果，使用Streamlit框架构建Web交互应用，实现用户输入消费画像后自动输出消费群体标签、MBTI性格预测和TOP-5搭子推荐。', style='List Bullet')

doc.add_heading('1.3 技术栈', level=2)
doc.add_paragraph(
    '编程语言：Python 3.x\n'
    '数据处理：Pandas, NumPy\n'
    '机器学习：Scikit-learn（K-Means, Random Forest, StandardScaler, LOOCV）\n'
    '不平衡处理：Imbalanced-learn（SMOTENC）\n'
    '数据可视化：Matplotlib, Seaborn, Plotly\n'
    'Web框架：Streamlit\n'
    '开发环境：VS Code + Anaconda (ML_Course)'
)

doc.add_page_break()

# ========== 第二章：数据采集与预处理 ==========
doc.add_heading('二、数据采集与预处理', level=1)

doc.add_heading('2.1 问卷设计', level=2)
doc.add_paragraph(
    '问卷共包含19道题目，分为五个模块：\n'
    '（1）基础信息（4题）：性别、年级、生活费来源、月生活费区间\n'
    '（2）细分消费结构（5题）：餐饮伙食、社交娱乐、服饰美妆、学习发展、数码游戏的每月花费\n'
    '（3）消费行为与认知（3题）：消费决策习惯、选购看重因素、兴趣投入意愿\n'
    '（4）MBTI性格快速自测（4题）：分别对应E/I、S/N、T/F、J/P四个维度\n'
    '（5）交友与场景推荐意向（3题）：消费观念朋友偏好、推荐系统态度、心仪结伴场景\n\n'
    '问卷通过问卷星平台发放，最终回收有效答卷69份。'
    '其中第13题（E/I维度）因平台设置有误额外包含"两者皆可"选项，'
    '导致该维度成为三分类问题（I/E/X），其余三个维度保持二分类。'
)

doc.add_heading('2.2 数据清洗', level=2)
doc.add_paragraph(
    '对原始69条数据执行以下清洗规则：'
)
doc.add_paragraph('规则一：填写时间短于30秒的答卷标记为无效（剔除4份）。30秒是阅读并理解全部19道题的最低时间下限。', style='List Bullet')
doc.add_paragraph('规则二：五个消费维度答案完全相同的答卷标记为无效（剔除5份）。此类答卷通常为"全部选同一选项"的应付式填写，不反映真实消费差异。', style='List Bullet')
doc.add_paragraph(
    '注：阶段一初版脚本曾包含第三条规则"MBTI四维答案全同标记为无效"，'
    '经逐条审查发现被标记的11条记录全部为真实ISTJ类型（4题均选第1选项，恰好构成ISTJ），'
    '其中真正乱填者已被前两条规则覆盖（如填写时间24秒且消费全同的#9号答卷）。'
    '因此最终版本移除该规则，避免误伤真实样本。\n\n'
    '清洗后保留有效数据62份（剔除率10.1%）。'
)

doc.add_heading('2.3 数据编码', level=2)
doc.add_paragraph(
    '问卷选项均以数值编码存储，映射关系如下表所示：'
)

# 编码表
table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
hdr = table.rows[0].cells
hdr[0].text = '变量'
hdr[1].text = '编码'
hdr[2].text = '含义'

encodings = [
    ('gender', '1 / 2', '男 / 女'),
    ('monthly_budget', '1~5', '≤1000 / 1001-1500 / 1501-2000 / 2001-3000 / >3000元'),
    ('consume_food', '1~4', '<600 / 600-900 / 901-1200 / >1200元'),
    ('consume_social', '1~4', '<100 / 100-300 / 301-500 / >500元'),
    ('consume_fashion', '1~4', '几乎不花 / 100-300 / 301-500 / >500元'),
    ('consume_study', '1~4', '几乎不花 / <50 / 50-150 / >150元'),
    ('consume_digital', '1~4', '几乎不花 / <100 / 100-300 / >300元'),
    ('mbti_ei', '1 / 2 / 3', 'I(内倾) / E(外倾) / X(两者皆可)'),
    ('mbti_sn/jp/tf', '1 / 2', 'S/N/T/J / N/F/P 等'),
]
for var, code, meaning in encodings:
    row = table.add_row().cells
    row[0].text = var
    row[1].text = code
    row[2].text = meaning

doc.add_paragraph()
doc.add_paragraph('[此处插入代码: phase1_clustering.py 第57-82行 — LABEL_MAP 标签映射字典]')

doc.add_page_break()

# ========== 第三章：阶段一 —— 无监督聚类 ==========
doc.add_heading('三、阶段一：消费群体无监督聚类', level=1)

doc.add_heading('3.1 特征选择与标准化', level=2)
doc.add_paragraph(
    '选取7个特征作为聚类输入：五项消费明细（consume_food, consume_social, consume_fashion, '
    'consume_study, consume_digital）、月生活费区间（monthly_budget）、兴趣投入意愿（hobby_spend）。\n'
    '所有特征值域为1~4或1~5的离散等级，使用StandardScaler进行Z-score标准化，消除量纲差异。'
)

doc.add_heading('3.2 最优K值确定', level=2)
doc.add_paragraph(
    '采用肘部法则（Elbow Method）与轮廓系数（Silhouette Score）双重准则确定最优聚类数K。\n'
    '在K=2~8范围内遍历，计算每个K值对应的簇内平方和（Inertia）和轮廓系数。'
    '结果显示，K=4时Inertia下降速率出现明显拐点（ΔInertia从33.0骤降至14.7），'
    '且轮廓系数为0.1775（在K≥3中较高），综合判定K=4为最优聚类数。'
)

doc.add_paragraph('[此处插入图片: output_elbow_silhouette.png — 肘部法则与轮廓系数图]')

doc.add_heading('3.3 聚类结果与群体画像', level=2)
doc.add_paragraph(
    'K-Means（K=4, random_state=42, n_init=10）聚类结果将62份样本分为四个群体，'
    '各群体消费特征均值如下表所示：'
)

# 聚类汇总表
table2 = doc.add_table(rows=1, cols=8, style='Light Grid Accent 1')
hdr2 = table2.rows[0].cells
hdr2[0].text = '群体'
hdr2[1].text = '人数'
hdr2[2].text = '月预算'
hdr2[3].text = '餐饮'
hdr2[4].text = '社交'
hdr2[5].text = '服饰'
hdr2[6].text = '学习'
hdr2[7].text = '数码'

cluster_data = [
    ('0 高消费数码型', '12', '4.08', '3.50', '2.58', '1.67', '1.67', '3.17'),
    ('1 节俭生存型', '23', '2.13', '2.04', '1.43', '1.13', '1.70', '2.17'),
    ('2 社交体验型', '14', '3.07', '2.57', '2.64', '2.21', '1.57', '1.79'),
    ('3 自我提升型', '13', '3.15', '2.69', '2.46', '2.77', '3.00', '2.08'),
]
for cd in cluster_data:
    row = table2.add_row().cells
    for i, val in enumerate(cd):
        row[i].text = val

doc.add_paragraph()
doc.add_paragraph('四个消费群体的特征解读如下：')
doc.add_paragraph(
    '群体0 — 高消费数码型（19.4%）：月预算全组最高（2001-3000元），餐饮和数码消费双高，'
    '但服饰和学习投入极低。典型画像：生活费和伙食费充裕，大量投入游戏和电子设备，不太注重穿着打扮。'
    '主流MBTI为XNTP。', style='List Bullet'
)
doc.add_paragraph(
    '群体1 — 节俭生存型（37.1%）：五项消费维度全部垫底，月预算集中在1001-1500元。'
    '典型画像：消费极端克制，仅覆盖基本餐饮需求，在社交、服饰等方面几乎零花销。'
    '主流MBTI为ISFP。', style='List Bullet'
)
doc.add_paragraph(
    '群体2 — 社交体验型（22.6%）：社交娱乐支出突出（2.64），但学习和数码投入极低（1.57、1.79）。'
    '典型画像："重人际、轻资产"，愿意为聚餐、出游等社交活动买单，但不爱学习也不打游戏。'
    '主流MBTI为XSFP。', style='List Bullet'
)
doc.add_paragraph(
    '群体3 — 自我提升型（21.0%）：学习发展和服饰美妆均为全组最高（3.00、2.77），兴趣投入意愿也领先（2.38）。'
    '典型画像：内外兼修，消费带有明显的"投资自己"属性，既注重外貌也注重知识提升。'
    '主流MBTI为XNFP。', style='List Bullet'
)

doc.add_paragraph('[此处插入图片: output_radar.png — 四群体消费雷达图]')
doc.add_paragraph('[此处插入图片: output_heatmap.png — 群体×消费特征热力图]')
doc.add_paragraph('[此处插入图片: output_pca.png — PCA二维聚类分布散点图]')

doc.add_paragraph('[此处插入代码: phase1_clustering.py 第135-180行 — K-Means聚类核心代码段]')

doc.add_page_break()

# ========== 第四章：阶段二 —— MBTI预测模型 ==========
doc.add_heading('四、阶段二：MBTI性格预测模型', level=1)

doc.add_heading('4.1 任务定义', level=2)
doc.add_paragraph(
    '阶段二的目标是验证"消费行为可以预测MBTI人格"的假设。具体而言：\n'
    '输入（X）：12维特征，包括5个消费维度、4个消费行为/人口学特征（月预算、兴趣投入、决策习惯、看重因素）、3个人口学特征（性别、年级、生活费来源）\n'
    '输出（Y）：4个独立目标——E/I（三分类）、S/N（二分类）、T/F（二分类）、J/P（二分类）\n'
    '模型结构：1个三分类器 + 3个二分类器，各维度独立训练与评估。'
)

doc.add_heading('4.2 小样本应对策略', level=2)
doc.add_paragraph(
    '62份样本在机器学习中属于极小样本量，常规7:3训练-测试划分将导致测试集仅18条，'
    '评估结果方差极大。为此采用以下三项针对性策略：'
)
doc.add_paragraph(
    '留一法交叉验证（LOOCV）：每次以61条样本训练、1条样本测试，重复62次取平均。'
    'LOOCV在小样本下提供几乎无偏的泛化误差估计，是最适合本场景的验证方法。', style='List Bullet'
)
doc.add_paragraph(
    'SMOTENC过采样：针对E/I维度中外倾（E）样本仅9份（14.5%）的严重不平衡问题，'
    '在每折LOOCV的训练集上应用SMOTENC（Synthetic Minority Over-sampling Technique for '
    'Nominal and Continuous features）合成少数类样本。SMOTENC能够正确处理决策习惯和性别等类别型特征，'
    '避免传统SMOTE对离散变量的不当插值。', style='List Bullet'
)
doc.add_paragraph(
    '模型正则化：Random Forest设置max_depth=5、min_samples_leaf=3，'
    '配合class_weight="balanced"，在控制过拟合的同时缓解类别不平衡。', style='List Bullet'
)

doc.add_heading('4.3 模型性能', level=2)
doc.add_paragraph(
    '四个维度的LOOCV评估结果汇总如下：'
)

table3 = doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
hdr3 = table3.rows[0].cells
hdr3[0].text = '维度'
hdr3[1].text = '类别数'
hdr3[2].text = 'LOOCV准确率'
hdr3[3].text = '随机基线'
hdr3[4].text = '提升幅度'

perf_data = [
    ('E/I（精力恢复）', '3', '48.4%', '33.3%', '+15.1%'),
    ('S/N（信息处理）', '2', '61.3%', '50.0%', '+11.3%'),
    ('T/F（决策标准）', '2', '45.2%', '50.0%', '-4.8%'),
    ('J/P（生活风格）', '2', '58.1%', '50.0%', '+8.1%'),
]
for pd_row in perf_data:
    row = table3.add_row().cells
    for i, val in enumerate(pd_row):
        row[i].text = val

doc.add_paragraph()
doc.add_paragraph(
    '四维平均准确率为53.2%。其中S/N维度表现最佳（61.3%），表明消费特征对"感觉/直觉"信息处理方式具有一定预测力。'
    'J/P维度次之（58.1%），消费行为中的决策习惯和看重因素可能反映了一个人的计划性倾向。'
    'E/I维度准确率48.4%虽高于随机基线（33.3%），但E类（外倾）的召回率为0——模型完全未能习得外倾样本的消费模式，'
    '这与E类仅9份样本的客观限制一致。T/F维度准确率45.2%低于随机基线，表明消费行为与"理性/感性"决策偏好之间关联较弱。'
)

doc.add_paragraph('[此处插入图片: output_confusion_matrices.png — 四维混淆矩阵]')

doc.add_heading('4.4 特征重要性分析', level=2)
doc.add_paragraph(
    '使用全部62份样本训练的全量模型提取随机森林特征重要性（Gini Importance），各维度Top-3特征如下：'
)
doc.add_paragraph('E/I维度：性别(0.135) > 数码游戏(0.122) > 社交娱乐(0.121)。社交娱乐和数码游戏分别排名第3和第2，方向性符合"社交→E、数码→I"的初始假设，但性别以微弱优势居首，提示该维度可能更受人口学变量影响。', style='List Bullet')
doc.add_paragraph('S/N维度：社交娱乐(0.218) > 学习发展(0.099) > 餐饮伙食(0.088)。社交娱乐以显著优势领先，直觉型(N)可能更倾向于多样化的社交体验和知识探索。', style='List Bullet')
doc.add_paragraph('T/F维度：服饰美妆(0.191) > 社交娱乐(0.177) > 餐饮伙食(0.115)。服饰美妆消费对"理性/感性"维度的影响最大，可能反映了情感型(F)更注重外在形象。', style='List Bullet')
doc.add_paragraph('J/P维度：看重因素(0.167) > 决策习惯(0.131) > 数码游戏(0.110)。消费决策中的"看重什么"和"如何决策"直接反映了一个人的计划性，与理论预期吻合。', style='List Bullet')

doc.add_paragraph('[此处插入图片: output_feature_importance.png — 四维特征重要性条形图]')
doc.add_paragraph('[此处插入代码: phase2_mbti_classifier.py 第110-155行 — LOOCV+SMOTENC训练循环]')

doc.add_page_break()

# ========== 第五章：阶段三 —— 推荐系统与Web展示 ==========
doc.add_heading('五、阶段三：校园搭子推荐系统', level=1)

doc.add_heading('5.1 系统架构', level=2)
doc.add_paragraph(
    '阶段三使用Streamlit框架将阶段一和阶段二的模型封装为交互式Web应用。'
    '启动时加载预训练的K-Means聚类器（model_kmeans.pkl）、4个随机森林分类器（model_rf_*.pkl）'
    '以及对应的StandardScaler标准化器。'
    '用户通过侧边栏输入个人消费画像（7个滑块/下拉框）和可选的MBTI快速自测（4道选择题），'
    '点击"开始分析"后系统执行以下流程：'
)
doc.add_paragraph('用户输入 → StandardScaler标准化 → K-Means聚类（输出消费群体标签）', style='List Number')
doc.add_paragraph('用户输入（补全12维） → 4个RF分类器并行预测 → 拼接MBTI类型', style='List Number')
doc.add_paragraph('综合打分：消费余弦相似度(40%) + 预算匹配(20%) + MBTI互补(20%) + 场景契合(20%) → TOP-5推荐', style='List Number')

doc.add_heading('5.2 推荐算法', level=2)
doc.add_paragraph(
    '推荐引擎对62份真实样本逐一计算综合匹配分。四个评分维度如下：'
)
doc.add_paragraph('消费相似度（权重40%）：计算用户与候选人的7维消费特征向量的余弦相似度，归一化到[0,1]。余弦相似度关注消费结构的"形状"而非绝对金额，更能反映消费偏好的一致性。', style='List Bullet')
doc.add_paragraph('预算匹配（权重20%）：月生活费区间差值≤1档为满分（1.0），差2档为半满分（0.5），差3档及以上为0。控制绝对消费水平偏差。', style='List Bullet')
doc.add_paragraph('MBTI互补（权重20%）：依据经典MBTI配对理论（如ENFP×INTJ、ESFP×ISTJ等8对），若用户与候选人为互补型则得满分1.0，否则给予基线分0.3。', style='List Bullet')
doc.add_paragraph('场景契合（权重20%）：两人在"最心仪结伴场景"上选择相同则得1.0，否则为0。', style='List Bullet')

doc.add_heading('5.3 界面设计', level=2)
doc.add_paragraph(
    'Web应用分为三个标签页：\n'
    'Tab1「消费群体分析」：展示用户所属消费群体名称，以及用户消费雷达图与群体均值的对比（Plotly交互图）。\n'
    'Tab2「MBTI性格预测」：以4个Metric卡片展示四个维度的预测结果，下方显示综合MBTI类型。\n'
    'Tab3「最佳搭子推荐」：卡片式展示TOP-5匹配搭子，每张卡片包含MBTI类型、消费群体、偏好场景、四维匹配分条形图。'
)

doc.add_paragraph('[此处插入图片: phase3界面截图1 — 消费群体分析Tab]')
doc.add_paragraph('[此处插入图片: phase3界面截图2 — MBTI预测Tab]')
doc.add_paragraph('[此处插入图片: phase3界面截图3 — 搭子推荐Tab]')
doc.add_paragraph('[此处插入代码: phase3_apptest.py — Streamlit完整代码]')

doc.add_page_break()

# ========== 第六章：总结与展望 ==========
doc.add_heading('六、总结与展望', level=1)

doc.add_heading('6.1 项目成果总结', level=2)
doc.add_paragraph(
    '本项目完整实现了从数据采集、清洗、聚类分析、分类预测到推荐系统部署的完整机器学习流程：\n'
    '（1）基于62份真实问卷数据，通过K-Means聚类成功识别出高消费数码型、节俭生存型、社交体验型、自我提升型四类校园消费群体，群体间特征差异显著。\n'
    '（2）采用LOOCV+SMOTENC策略训练随机森林分类器，在S/N维度取得61.3%的预测准确率，验证了"消费行为可以部分预测MBTI人格"的核心假设。\n'
    '（3）基于Streamlit构建了可交互的Web推荐系统，实现了从消费画像输入到搭子匹配的端到端流程。'
)

doc.add_heading('6.2 局限性', level=2)
doc.add_paragraph('（1）样本量瓶颈：62份样本在机器学习中属于极小样本，尤其E/I维度仅9份外倾样本，直接导致模型对外倾类的预测完全失效。扩大问卷发放范围是提升模型鲁棒性的首要任务。', style='List Bullet')
doc.add_paragraph('（2）MBTI测量简化：问卷中每个MBTI维度仅用1道题测量，精度远低于标准93题MBTI量表。自评结果可能受社会期望偏差影响。', style='List Bullet')
doc.add_paragraph('（3）推荐冷启动：当前推荐范围仅限于62份已知样本。若系统面向更大用户群，需考虑新用户冷启动问题。', style='List Bullet')
doc.add_paragraph('（4）E/I维度三分类：因问卷平台设置失误导致E/I维度出现"两者皆可"选项，使得该维度与其他三个维度在类别数上不一致，增加了建模复杂度。', style='List Bullet')

doc.add_heading('6.3 未来改进方向', level=2)
doc.add_paragraph('（1）数据增强：采用SDV GaussianCopulaSynthesizer合成虚拟样本，将有效样本扩充至200-300条，提升聚类稳定性和分类模型泛化能力。', style='List Bullet')
doc.add_paragraph('（2）特征工程：引入更多消费行为特征（如支付方式偏好、品牌忠诚度、促销敏感度等），增强特征空间对MBTI的解释力。', style='List Bullet')
doc.add_paragraph('（3）模型优化：尝试XGBoost、LightGBM等集成模型，或通过网格搜索调优随机森林超参数。', style='List Bullet')
doc.add_paragraph('（4）推荐算法深化：引入协同过滤或基于内容的推荐，结合用户反馈数据持续优化匹配权重。', style='List Bullet')

doc.add_page_break()

# ========== 附录 ==========
doc.add_heading('附录', level=1)

doc.add_heading('A. 文件清单', level=2)
files = [
    ('MBTI性格与大学生消费行为关联.xlsx', '原始问卷数据（69份）'),
    ('phase1_clustering.py', '阶段一：数据清洗与K-Means聚类脚本'),
    ('phase2_mbti_classifier.py', '阶段二：MBTI预测模型（RF+LOOCV+SMOTENC）'),
    ('phase3_apptest.py', '阶段三：Streamlit推荐系统Web应用'),
    ('data_cleaned.csv', '清洗后数据集（62份）'),
    ('cluster_summary.csv', '聚类结果汇总表'),
    ('model_kmeans.pkl / model_scaler.pkl', '阶段一训练模型文件'),
    ('model_rf_*.pkl / model_scaler_*.pkl', '阶段二训练模型文件（4+4个）'),
    ('output_elbow_silhouette.png', '肘部法则与轮廓系数图'),
    ('output_radar.png', '消费群体雷达图'),
    ('output_heatmap.png', '群体特征热力图'),
    ('output_pca.png', 'PCA聚类分布散点图'),
    ('output_confusion_matrices.png', '四维混淆矩阵'),
    ('output_feature_importance.png', '四维特征重要性条形图'),
]
for fname, desc in files:
    doc.add_paragraph(f'{fname} — {desc}', style='List Bullet')

doc.add_heading('B. 参考文献', level=2)
doc.add_paragraph('[此处根据学校要求格式补充参考文献，建议包含：]')
doc.add_paragraph('K-Means算法原始论文：MacQueen, J. (1967). Some methods for classification and analysis of multivariate observations.', style='List Bullet')
doc.add_paragraph('随机森林：Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32.', style='List Bullet')
doc.add_paragraph('SMOTE过采样：Chawla, N. V., et al. (2002). SMOTE: Synthetic minority over-sampling technique. JAIR, 16, 321-357.', style='List Bullet')
doc.add_paragraph('LOOCV交叉验证：Stone, M. (1974). Cross-validatory choice and assessment of statistical predictions. JRSSB, 36(2), 111-133.', style='List Bullet')
doc.add_paragraph('SDV合成数据：Patki, N., et al. (2016). The Synthetic Data Vault. IEEE DSAA.', style='List Bullet')

# ========== 保存 ==========
doc.save('项目报告_校园圈层聚类与推荐系统.docx')
print('✅ 报告已生成: 项目报告_校园圈层聚类与推荐系统.docx')
